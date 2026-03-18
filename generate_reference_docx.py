"""
DocuVault Client Reference Guide — DOCX Generator
Theme: Renata IoT Reference Guide
Colors: H1=#E84C25, H2=#2C3E50, Accent=#4472C4
Logo: Logo.png
"""

import os, datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

LOGO_PATH = r'D:\AI_Model_Renata\Document-management\Group\V2\DocuVault\Logo.png'

# ── Colors ────────────────────────────────────────────────────────────────────
H1_COLOR   = RGBColor(0xE8, 0x4C, 0x25)   # orange-red
H2_COLOR   = RGBColor(0x2C, 0x3E, 0x50)   # dark navy
H3_COLOR   = RGBColor(0x2E, 0x74, 0xB5)   # blue
ACCENT     = RGBColor(0x44, 0x72, 0xC4)   # primary blue
ORANGE     = RGBColor(0xED, 0x7D, 0x31)   # orange
NAVY       = RGBColor(0x44, 0x54, 0x6A)   # dark table header
MID_GRAY   = RGBColor(0x6B, 0x72, 0x80)
BODY_TEXT  = RGBColor(0x2C, 0x2C, 0x2C)
WHITE_RGB  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG   = RGBColor(0xF8, 0xF9, 0xFA)
BORDER_COL = RGBColor(0xD1, 0xD5, 0xDB)

# ── Hex helpers ───────────────────────────────────────────────────────────────
def rgb_hex(r, g, b):
    return f'{r:02X}{g:02X}{b:02X}'

H1_HEX  = 'E84C25'
H2_HEX  = '2C3E50'
H3_HEX  = '2E74B5'
ACC_HEX = '4472C4'
NAV_HEX = '44546A'
WH_HEX  = 'FFFFFF'
LB_HEX  = 'F8F9FA'
BD_HEX  = 'D1D5DB'
OR_HEX  = 'FFF3ED'


# ── XML Helpers ───────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   val.get('val', 'single'))
            el.set(qn('w:sz'),    val.get('sz', '4'))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', 'auto'))
            borders.append(el)
    tcPr.append(borders)


def set_para_border_bottom(para, hex_color='E84C25', sz='12'):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    sz)
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), hex_color)
    pBdr.append(bot)
    pPr.append(pBdr)


def set_para_shading(para, hex_fill, hex_color='auto'):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), hex_color)
    shd.set(qn('w:fill'),  hex_fill)
    pPr.append(shd)


def set_run_highlight(run, hex_color):
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    rPr.append(shd)


def add_page_break(doc):
    doc.add_page_break()


def set_table_style(table):
    """Remove default table style borders, we set them manually."""
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    # Remove existing borders
    for b in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(b)
    borders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), BD_HEX)
        borders.append(el)
    tblPr.append(borders)


def set_col_widths(table, widths_cm):
    """Set column widths."""
    tbl = table._tbl
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is None:
        tblGrid = OxmlElement('w:tblGrid')
        tbl.insert(0, tblGrid)
    else:
        for col in tblGrid.findall(qn('w:gridCol')):
            tblGrid.remove(col)
    for w in widths_cm:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(int(w * 567)))   # 1 cm ≈ 567 twips
        tblGrid.append(gc)


# ── Document Setup ────────────────────────────────────────────────────────────

def setup_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.2)
    sec.right_margin  = Cm(2.2)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.0)

    # Default paragraph style
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10)
    normal.font.color.rgb = BODY_TEXT

    return doc


# ── Style Helpers ─────────────────────────────────────────────────────────────

def add_h1(doc, text):
    """H1 — orange-red, 16pt bold, with orange bottom border."""
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold       = True
    run.font.size  = Pt(16)
    run.font.color.rgb = H1_COLOR
    run.font.name  = 'Calibri'
    set_para_border_bottom(p, H1_HEX, '12')
    return p


def add_h2(doc, text):
    """H2 — dark navy, 12pt bold."""
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold       = True
    run.font.size  = Pt(12)
    run.font.color.rgb = H2_COLOR
    run.font.name  = 'Calibri'
    return p


def add_h3(doc, text):
    """H3 — blue, 11pt bold."""
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold       = True
    run.font.size  = Pt(11)
    run.font.color.rgb = H3_COLOR
    run.font.name  = 'Calibri'
    return p


def add_body(doc, text, bold_prefix=''):
    """Normal body paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        r = p.add_run(bold_prefix + ' ')
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = BODY_TEXT
        r.font.name = 'Calibri'
    run = p.add_run(text)
    run.font.size  = Pt(10)
    run.font.color.rgb = BODY_TEXT
    run.font.name  = 'Calibri'
    return p


def add_bullet(doc, text, bold_prefix='', level=0):
    """Bullet point paragraph."""
    p   = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.6 + level * 0.5)
    if bold_prefix:
        r = p.add_run(bold_prefix + ': ')
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = H2_COLOR
        r.font.name = 'Calibri'
    run = p.add_run(text)
    run.font.size  = Pt(10)
    run.font.color.rgb = BODY_TEXT
    run.font.name  = 'Calibri'
    return p


def add_note(doc, text):
    """Italic note paragraph with light blue background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    set_para_shading(p, 'EFF6FF')
    run = p.add_run('Note: ' + text)
    run.italic     = True
    run.font.size  = Pt(9.5)
    run.font.color.rgb = H3_COLOR
    run.font.name  = 'Calibri'
    return p


def add_info_box(doc, title, body_text, fill_hex='EFF6FF'):
    """Shaded info box with bold title and body."""
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(6)
    p1.paragraph_format.space_after  = Pt(2)
    p1.paragraph_format.left_indent  = Cm(0.3)
    set_para_shading(p1, fill_hex)
    r = p1.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = H2_COLOR
    r.font.name = 'Calibri'

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(8)
    p2.paragraph_format.left_indent  = Cm(0.3)
    p2.paragraph_format.right_indent = Cm(0.3)
    set_para_shading(p2, fill_hex)
    r2 = p2.add_run(body_text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = BODY_TEXT
    r2.font.name = 'Calibri'


def add_spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(pts)
    p.paragraph_format.space_before = Pt(0)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = MID_GRAY
    run.font.name = 'Calibri'


# ── Table Builder ─────────────────────────────────────────────────────────────

def add_table(doc, rows, col_widths_cm, header_hex=None):
    """
    rows: list of lists of strings. rows[0] = header.
    col_widths_cm: list of column widths in cm.
    """
    hcol = header_hex or H2_HEX
    n_cols = len(rows[0])
    table  = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_style(table)
    set_col_widths(table, col_widths_cm)

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        row.height = Cm(0.75)
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            p.paragraph_format.left_indent  = Cm(0.15)

            # Handle multiline cell text
            lines = str(cell_text).split('\n')
            for li, line in enumerate(lines):
                if li > 0:
                    p.add_run('\n')
                run = p.add_run(line)
                run.font.name = 'Calibri'
                if i == 0:
                    run.bold = True
                    run.font.size  = Pt(9.5)
                    run.font.color.rgb = WHITE_RGB
                    set_cell_bg(cell, hcol)
                else:
                    run.font.size  = Pt(9)
                    run.font.color.rgb = BODY_TEXT
                    if i % 2 == 0:
                        set_cell_bg(cell, LB_HEX)
                    else:
                        set_cell_bg(cell, WH_HEX)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


# ── Flow Diagram (text-based table) ──────────────────────────────────────────

def add_flow_diagram(doc, steps, colors_hex, label):
    """
    Render a horizontal flow as a single-row table with arrows.
    steps: list of (title, subtitle) tuples
    colors_hex: list of hex strings for each step box
    """
    n = len(steps)
    # Build a single-row table: step | arrow | step | arrow | ...
    n_cols = n * 2 - 1
    table  = doc.add_table(rows=1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Remove all table borders
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for b in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(b)
    bdr = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        bdr.append(el)
    tblPr.append(bdr)

    # Page width minus margins ≈ 16.6 cm
    avail  = 16.6
    bw     = (avail - (n - 1) * 0.8) / n
    row    = table.rows[0]
    row.height = Cm(1.3)

    for i, (title, sub) in enumerate(steps):
        ci   = i * 2
        cell = row.cells[ci]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        col  = colors_hex[i] if i < len(colors_hex) else ACC_HEX
        set_cell_bg(cell, col)

        # Remove cell borders
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        cb   = OxmlElement('w:tcBorders')
        for side in ['top','left','bottom','right']:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'none')
            cb.append(el)
        tcPr.append(cb)

        # Set cell width
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'),    str(int(bw * 567)))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(1)
        r1 = p.add_run(title)
        r1.bold = True
        r1.font.size = Pt(8)
        r1.font.color.rgb = WHITE_RGB
        r1.font.name = 'Calibri'
        if sub:
            r2 = p.add_run('\n' + sub)
            r2.font.size = Pt(7)
            r2.font.color.rgb = RGBColor(0xDB, 0xEA, 0xFE)
            r2.font.name = 'Calibri'

        # Arrow cell
        if i < n - 1:
            acell = row.cells[ci + 1]
            acell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_bg(acell, 'F8F9FA')

            atc   = acell._tc
            atcPr = atc.get_or_add_tcPr()
            acb   = OxmlElement('w:tcBorders')
            for side in ['top','left','bottom','right']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'none')
                acb.append(el)
            atcPr.append(acb)
            atcW = OxmlElement('w:tcW')
            atcW.set(qn('w:w'),    str(int(0.8 * 567)))
            atcW.set(qn('w:type'), 'dxa')
            atcPr.append(atcW)

            ap = acell.paragraphs[0]
            ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ar = ap.add_run('\u2192')    # →
            ar.font.size = Pt(14)
            ar.font.color.rgb = H2_COLOR
            ar.font.name = 'Calibri'

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    add_caption(doc, label)


# ── Centered image helper ─────────────────────────────────────────────────────

def add_image_centered(doc, img_path, width_cm, caption_text=''):
    """Insert an image centered on the page with optional caption."""
    if not os.path.exists(img_path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run()
    run.add_picture(img_path, width=Cm(width_cm))
    if caption_text:
        add_caption(doc, caption_text)


# ── Arch diagram (text table) ─────────────────────────────────────────────────

def add_arch_table(doc):
    """Three-layer architecture as a formatted table."""
    data = [
        ['Layer',              'What it is',                                 'What it does for you'],
        ['Your Browser',       'The web interface — open on any device.',    'Where you upload, search, chat with AI, and manage settings.'],
        ['DocuVault Platform', 'The server application behind the scenes.',  'Handles login, stores documents, controls who sees what, and connects to AI.'],
        ['AI & Storage',       'The AI engine and databases.',               'Reads and indexes documents, answers questions, and stores all data securely.'],
    ]
    add_table(doc, data, [3.5, 5.5, 7.6], H2_HEX)


# ── Header / Footer ───────────────────────────────────────────────────────────

def add_header_footer(doc):
    section = doc.sections[0]

    # ── Header ──
    header = section.header
    header.is_linked_to_previous = False
    # Clear default content
    for p in header.paragraphs:
        for run in p.runs:
            run.text = ''

    htable = header.add_table(1, 2, Cm(16.6))
    htable.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Remove table borders
    htbl  = htable._tbl
    htblP = htbl.find(qn('w:tblPr'))
    if htblP is None:
        htblP = OxmlElement('w:tblPr')
        htbl.insert(0, htblP)
    hbdr = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        hbdr.append(el)
    htblP.append(hbdr)

    hrow  = htable.rows[0]
    hrow.height = Cm(0.9)

    # Left cell — logo
    lcell = hrow.cells[0]
    lcell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_bg(lcell, H2_HEX)
    lp = lcell.paragraphs[0]
    lp.paragraph_format.left_indent = Cm(0.2)
    if os.path.exists(LOGO_PATH):
        lrun = lp.add_run()
        lrun.add_picture(LOGO_PATH, width=Cm(3.5))
    else:
        lr = lp.add_run('Renata AI')
        lr.bold = True
        lr.font.size = Pt(11)
        lr.font.color.rgb = WHITE_RGB
        lr.font.name = 'Calibri'

    # Right cell — doc title
    rcell = hrow.cells[1]
    rcell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_bg(rcell, H2_HEX)
    rp = rcell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.right_indent = Cm(0.3)
    rr = rp.add_run('DocuVault  |  Client Reference Guide')
    rr.font.size = Pt(8.5)
    rr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    rr.font.name = 'Calibri'

    # Orange accent line below header
    hp = header.add_paragraph()
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after  = Pt(0)
    set_para_border_bottom(hp, H1_HEX, '8')

    # ── Footer ──
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        for run in p.runs:
            run.text = ''

    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.paragraph_format.space_before = Pt(4)
    fp.paragraph_format.space_after  = Pt(0)
    set_para_border_bottom(fp, 'E2E8F0', '4')

    fl = fp.add_run(
        f'Renata Envirocom Pvt. Ltd.  |  Confidential  |  '
        f'{datetime.date.today().strftime("%B %Y")}')
    fl.font.size = Pt(8)
    fl.font.color.rgb = MID_GRAY
    fl.font.name = 'Calibri'

    # Page number (right side via tab)
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pr = fp.add_run('\t\tPage ')
    pr.font.size = Pt(8)
    pr.font.color.rgb = MID_GRAY
    pr.font.name = 'Calibri'

    # Insert PAGE field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run_el = OxmlElement('w:r')
    run_el.append(fldChar1)
    run_el.append(instrText)
    run_el.append(fldChar2)
    fp._p.append(run_el)

    # Tab stops for right-aligned page number
    pPr = fp._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9350')
    tabs.append(tab)
    pPr.append(tabs)


# ── Cover Page ────────────────────────────────────────────────────────────────

def build_cover(doc):
    # Orange header block
    p_logo = doc.add_paragraph()
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after  = Pt(0)
    set_para_shading(p_logo, H1_HEX)
    if os.path.exists(LOGO_PATH):
        run = p_logo.add_run()
        run.add_picture(LOGO_PATH, width=Cm(7))
    else:
        r = p_logo.add_run('Renata AI')
        r.bold = True
        r.font.size = Pt(24)
        r.font.color.rgb = WHITE_RGB
        r.font.name = 'Calibri'

    # Orange tagline bar
    p_tag = doc.add_paragraph()
    p_tag.paragraph_format.space_before = Pt(0)
    p_tag.paragraph_format.space_after  = Pt(0)
    set_para_shading(p_tag, H1_HEX)
    rt = p_tag.add_run('IoT  ·  AI  ·  Automation')
    rt.font.size = Pt(10)
    rt.font.color.rgb = RGBColor(0xFE, 0xCA, 0xCA)
    rt.font.name = 'Calibri'

    # Spacer
    for _ in range(4):
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Title
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after  = Pt(6)
    r1 = p1.add_run('DocuVault')
    r1.bold = True
    r1.font.size = Pt(36)
    r1.font.color.rgb = H1_COLOR
    r1.font.name = 'Calibri'

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run('Knowledge Management AI Platform')
    r2.bold = True
    r2.font.size = Pt(16)
    r2.font.color.rgb = H2_COLOR
    r2.font.name = 'Calibri'

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(16)
    r3 = p3.add_run('Client Reference Guide')
    r3.font.size = Pt(13)
    r3.font.color.rgb = MID_GRAY
    r3.font.name = 'Calibri'

    # Orange divider
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(20)
    set_para_border_bottom(p_div, H1_HEX, '16')

    # Meta table
    meta = [
        ['Version',         'DocuVault v2.0'],
        ['Date',            datetime.date.today().strftime('%B %d, %Y')],
        ['Prepared by',     'Renata AI'],
        ['Classification',  'Confidential – Client Use Only'],
    ]
    mt = doc.add_table(rows=len(meta), cols=2)
    mt.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_col_widths(mt, [3.5, 10])
    tbl = mt._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for b in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(b)
    bdr = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        bdr.append(el)
    tblPr.append(bdr)

    for row_data in meta:
        r_idx = meta.index(row_data)
        row   = mt.rows[r_idx]
        c0    = row.cells[0]
        c1    = row.cells[1]
        p0    = c0.paragraphs[0]
        p1    = c1.paragraphs[0]
        rl    = p0.add_run(row_data[0])
        rl.bold = True
        rl.font.size = Pt(9.5)
        rl.font.color.rgb = MID_GRAY
        rl.font.name = 'Calibri'
        rv = p1.add_run(row_data[1])
        rv.font.size = Pt(9.5)
        rv.font.color.rgb = BODY_TEXT
        rv.font.name = 'Calibri'
        for c in [c0, c1]:
            c.paragraphs[0].paragraph_format.space_before = Pt(3)
            c.paragraphs[0].paragraph_format.space_after  = Pt(3)

    # Footer area
    for _ in range(8):
        doc.add_paragraph().paragraph_format.space_after = Pt(14)

    p_co = doc.add_paragraph()
    r_co = p_co.add_run('Renata Envirocom Pvt. Ltd.  |  renataiot.com')
    r_co.font.size = Pt(9)
    r_co.font.color.rgb = MID_GRAY
    r_co.font.name = 'Calibri'

    doc.add_page_break()


# ── TOC ───────────────────────────────────────────────────────────────────────

def build_toc(doc):
    add_h1(doc, 'Contents')
    add_spacer(doc, 8)

    toc_items = [
        ('1.', 'System Overview', []),
        ('2.', 'Key Features', [
            ('2.1', 'Document Management'),
            ('2.2', 'AI Knowledge Assistant'),
            ('2.3', 'Users, Roles & Access Control'),
            ('2.4', 'Collaboration & Sharing'),
            ('2.5', 'Search & Organisation'),
            ('2.6', 'Audit Trail & Notifications'),
        ]),
        ('3.', 'Accessing System Features', []),
        ('4.', 'Architecture Overview', [
            ('4.1', 'How the System is Structured'),
            ('4.2', 'How Documents are Indexed'),
            ('4.3', 'How the AI Answers Questions'),
        ]),
        ('5.', 'Technical Reference', []),
        ('6.', 'Add-On Features & Integrations', [
            ('6.1', 'Annual Maintenance Contract (AMC)'),
            ('6.2', 'Email System Integration'),
            ('6.3', 'LAN Drive Sync (Google Drive-Style)'),
            ('6.4', 'Auto-Indexing on Upload'),
            ('6.5', 'Desktop Agent — Web to Desktop'),
            ('6.6', 'Offline Operation — No Internet Required'),
        ]),
        ('7.', 'Client Q&A Summary', []),
        ('8.', 'Support & Contact', []),
    ]

    for num, title, subs in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(f'{num}   {title}')
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = H2_COLOR
        r.font.name = 'Calibri'
        for snum, stitle in subs:
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after  = Pt(2)
            sp.paragraph_format.left_indent  = Cm(0.8)
            sr = sp.add_run(f'{snum}   {stitle}')
            sr.font.size = Pt(10)
            sr.font.color.rgb = MID_GRAY
            sr.font.name = 'Calibri'

    doc.add_page_break()


# ── 1. Overview ───────────────────────────────────────────────────────────────

def build_overview(doc):
    add_h1(doc, '1.  System Overview')
    add_spacer(doc, 6)

    add_body(doc,
        'DocuVault is Renata AI\'s Knowledge Management Platform. It gives your organisation '
        'one place to store, organise, and query all your documents — using plain English, not search keywords.')
    add_body(doc,
        'At its core is an AI assistant that reads your documents and answers questions instantly. '
        'No more hunting through folders or opening file after file.')

    add_h2(doc, 'What you can do with DocuVault')
    bullets = [
        'Upload and organise your documents in one secure place.',
        'Control who can see each document — from fully public to private or role-restricted.',
        'Ask the AI assistant questions in plain English and get answers drawn from your own documents.',
        'Collaborate with your team through comments, shared links, and notifications.',
        'Track every action with a complete audit trail.',
        'Manage users and roles with fine-grained permission levels.',
    ]
    for b in bullets:
        add_bullet(doc, b)

    add_spacer(doc, 8)
    add_table(doc, [
        ['Item',               'Details'],
        ['Platform',           'DocuVault v2.0 — web-based (browser) + optional desktop agent'],
        ['Supported Files',    'PDF, Word, text files, images, and more'],
        ['Max File Size',      '100 MB per upload'],
        ['AI Knowledge Modes', 'Hybrid (default)  ·  Strict (documents only)  ·  Indicated'],
        ['Access Control',     'Public  ·  Private  ·  Role-Based  ·  Custom (per user)'],
        ['Web Access',         'Any modern web browser on desktop, tablet, or mobile'],
        ['Desktop Access',     'Optional background desktop agent — auto-connects files, runs in system tray'],
        ['Internet Required',  'No — fully offline after one-time model download. No API key, no cloud dependency.'],
    ], [5.0, 11.6])
    doc.add_page_break()


# ── 2. Features ───────────────────────────────────────────────────────────────

def build_features(doc):
    add_h1(doc, '2.  Key Features')
    add_spacer(doc, 6)

    # 2.1
    add_h2(doc, '2.1  Document Management')
    add_body(doc, 'Everything you need to manage documents throughout their full lifecycle.')
    add_table(doc, [
        ['Feature',           'What it does'],
        ['Upload',            'Add documents from your computer. Supported types include PDF, Word, images, and more.'],
        ['Access Control',    'Choose who can see each document: Public, Private, Role-Based, or shared with specific people.'],
        ['Version History',   'Every time a document is edited, the previous version is saved. You can view or restore any past version.'],
        ['Document Locking',  'Lock a document while editing so no one else can make changes at the same time.'],
        ['Categories & Tags', 'Organise documents into categories (with sub-categories) and tag them for flexible grouping.'],
        ['Metadata',          'Track view count, download count, file size, and upload date automatically.'],
        ['Soft Delete',       'Deleted documents can be recovered by an administrator before permanent removal.'],
    ], [4.0, 12.6])
    add_spacer(doc, 8)

    # 2.2
    add_h2(doc, '2.2  AI Knowledge Assistant')
    add_body(doc,
        'The AI assistant lets you ask questions about your documents in plain English. '
        'It reads through all indexed documents, finds the most relevant parts, and writes a clear answer.')
    add_table(doc, [
        ['Feature',                'What it does'],
        ['Ask in Plain English',   'Type a question naturally — the AI understands context, not just keywords.'],
        ['Answers from Your Docs', 'The AI draws from documents stored in DocuVault, so answers are grounded in your data.'],
        ['Follow-up Questions',    'Ask follow-up questions in the same conversation. The AI remembers context.'],
        ['Three Answer Modes',
            'Hybrid (default): Uses your documents first, fills gaps with general knowledge.\n'
            'Strict: Only answers from your documents — refuses if the answer is not in any document.\n'
            'Indicated: Labels which parts come from your documents and which from general knowledge.'],
        ['Permission-Aware',       'The AI only uses documents the logged-in user is allowed to access.'],
        ['Source Transparency',    'Each answer shows which documents and pages were used, so you can verify the information.'],
    ], [4.0, 12.6])
    add_spacer(doc, 8)

    # 2.3
    add_h2(doc, '2.3  Users, Roles & Access Control')
    add_table(doc, [
        ['User Type',    'What they can do'],
        ['Guest',        'View documents that are set to Public only. Cannot upload or edit.'],
        ['Regular User', 'Upload and manage their own documents. Access documents based on their role level. Comment and collaborate.'],
        ['Admin',        'Full access to all documents, all users, and all system settings.'],
    ], [3.5, 13.1])
    add_spacer(doc, 4)
    add_body(doc,
        'Administrators can create custom roles (e.g. "Team Lead", "Department Head") and assign a numeric level to each. '
        'Documents set to Role-Based access are visible only to users whose role level meets or exceeds the requirement.')
    add_spacer(doc, 8)

    # 2.4
    add_h2(doc, '2.4  Collaboration & Sharing')
    add_table(doc, [
        ['Feature',         'What it does'],
        ['Comments',        'Add comments to any document. Replies are threaded for easy reading.'],
        ['Shared Links',    'Generate a link to share a document. Optionally add a password, set an expiry date, or limit how many times it can be opened.'],
        ['Direct Sharing',  'Share a document directly with specific registered users.'],
        ['Notifications',   'The system automatically alerts you when a document is shared with you, updated, or commented on.'],
        ['Favourites',      'Bookmark documents you use often for quick access from your sidebar.'],
    ], [3.5, 13.1])
    add_spacer(doc, 8)

    # 2.5
    add_h2(doc, '2.5  Search & Organisation')
    add_table(doc, [
        ['Feature',    'What it does'],
        ['Search Bar', 'Search by title, description, or content. Results rank the most relevant documents first.'],
        ['Filters',    'Narrow results by owner, date, access level, category, tags, or file type.'],
        ['Sorting',    'Sort by title, upload date, last updated, view count, or file size.'],
        ['Categories', 'Organise documents into a tree of categories. Each category can have a colour and icon.'],
        ['Tags',       'Add one or more tags to a document. Search and filter by tag across all categories.'],
    ], [3.5, 13.1])
    add_spacer(doc, 8)

    # 2.6
    add_h2(doc, '2.6  Audit Trail & Notifications')
    add_body(doc,
        'Every action in DocuVault is automatically recorded. The activity log shows who did what, '
        'and when — making it easy to track changes for compliance and accountability.')
    add_table(doc, [
        ['Recorded Action',     'When it appears in the log'],
        ['Upload / Create',     'A new document is added to the system.'],
        ['View / Download',     'A user opens or downloads a document.'],
        ['Edit / Update',       'A document or its details are changed.'],
        ['Delete',              'A document is removed (soft-deleted).'],
        ['Share',               'A document is shared via link or direct user share.'],
        ['Comment',             'A comment is added to a document.'],
        ['Permission Changed',  'A document\'s access level or user role is updated.'],
    ], [4.5, 12.1])
    doc.add_page_break()


# ── 3. Accessing Features ─────────────────────────────────────────────────────

def build_access(doc):
    add_h1(doc, '3.  Accessing System Features')
    add_spacer(doc, 6)
    add_body(doc,
        'The table below shows how to navigate to each feature. '
        'All paths are relative to your system\'s base URL (e.g. https://yourdomain.com).')
    add_spacer(doc, 6)

    rows = [
        ['URL Path',                   'How to access'],
        ['/register/',                 'Open the system in your browser and click Register. Enter your name, email, and password.'],
        ['/login/',                    'Enter your username and password and click Sign In.'],
        ['/dashboard/',                'Your home screen after login. Shows recent documents, notifications, and quick links.'],
        ['/documents/',                'The full document library. Use the search bar and filter panel to narrow results.'],
        ['/documents/create/',         'Click + New Document. Select a file, fill in the details, set access level, and click Save.'],
        ['/documents/<id>/',           'Click any document title to open it — preview, download, comment, version history, and share.'],
        ['/documents/<id>/edit/',      'Open a document then click Edit. Save changes — the system automatically creates a new version.'],
        ['/documents/<id>/index/',     'Open a document then click Index for AI. The system reads and indexes the content for the AI assistant.'],
        ['/documents/bulk-index/',     'Admins only. From the document list, go to Actions → Bulk Index to index all unprocessed documents.'],
        ['/chatbot/',                  'Click AI Assistant in the sidebar. Type your question and press Enter.'],
        ['/search/',                   'Use the search bar in the top navigation. Add filters for date, owner, category, tags, and access level.'],
        ['/categories/',               'Navigate to Organise → Categories to create, nest, and colour-code categories.'],
        ['/favorites/',                'Click the star icon on any document. Access bookmarks via Favourites in the sidebar.'],
        ['/notifications/',            'Click the bell icon in the top bar to see all notifications.'],
        ['/activity/',                 'Go to Account → Activity Log for a complete history of all actions.'],
        ['/admin/users/',              'Admins only. Go to Admin → Users to view all accounts and update role assignments.'],
        ['/admin/roles/',              'Admins only. Go to Admin → Roles to create, edit, or remove custom roles.'],
        ['/profile/edit/',             'Click your avatar (top right) → Edit Profile to update your details.'],
    ]

    n_cols = 2
    table  = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_style(table)
    set_col_widths(table, [4.5, 12.1])

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        row.height = Cm(0.75)
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            p.paragraph_format.left_indent  = Cm(0.15)

            if i == 0:
                run = p.add_run(cell_text)
                run.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = WHITE_RGB
                run.font.name = 'Calibri'
                set_cell_bg(cell, H2_HEX)
            else:
                if j == 0:
                    run = p.add_run(cell_text)
                    run.bold = True
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = ACCENT
                    run.font.name = 'Calibri'
                else:
                    run = p.add_run(cell_text)
                    run.font.size = Pt(9)
                    run.font.color.rgb = BODY_TEXT
                    run.font.name = 'Calibri'
                if i % 2 == 0:
                    set_cell_bg(cell, LB_HEX)
                else:
                    set_cell_bg(cell, WH_HEX)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    doc.add_page_break()


# ── 4. Architecture ───────────────────────────────────────────────────────────

def build_architecture(doc):
    add_h1(doc, '4.  Architecture Overview')
    add_spacer(doc, 6)
    add_body(doc,
        'This section gives a plain-English view of how DocuVault is structured — '
        'how it stores your documents, how the AI learns from them, and how it answers your questions. '
        'Every component described here runs on your own hardware with no internet connection required.')
    add_spacer(doc, 8)

    # 4.1
    add_h2(doc, '4.1  How the System is Structured')
    add_body(doc, 'DocuVault has three main layers that work together:')
    add_spacer(doc, 4)
    add_table(doc, [
        ['Layer',              'What it is',                                  'What it does for you'],
        ['Your Browser',       'The web interface you open on any device.',   'Where you upload, search, chat with AI, and manage settings.'],
        ['DocuVault Platform', 'The server application in the background.',   'Handles login, stores documents securely, controls access, and connects you to AI.'],
        ['AI & Storage',       'The AI engine and databases.',                'Reads and indexes documents, answers questions, and stores all data.'],
    ], [3.5, 5.5, 7.6])
    add_spacer(doc, 6)

    add_h3(doc, 'System Flow')
    add_flow_diagram(doc,
        [('Your Browser', 'Any device'),
         ('DocuVault\nPlatform', 'Application'),
         ('Document\nStorage', 'Local files'),
         ('AI Search\nEngine', 'Local index'),
         ('Ollama\n(Local AI)', 'No internet')],
        [H2_HEX, ACC_HEX, '1D4ED8', '059669', H1_HEX],
        'Figure 1: All components run on your local server — no internet required during operation.')
    add_spacer(doc, 6)

    # 4.2
    add_h2(doc, '4.2  How Documents are Indexed')
    add_body(doc,
        'Before the AI can answer questions about a document, it must be "indexed" — '
        'this means the AI reads it and saves a summary in a way it can search very quickly.')
    add_spacer(doc, 6)

    add_flow_diagram(doc,
        [('Upload\nDocument', ''),
         ('Platform\nValidates', ''),
         ('AI Reads &\nProcesses', ''),
         ('Saved to\nSearch Index', ''),
         ('Ready for\nAI Queries', '')],
        [ACC_HEX, H2_HEX, '059669', H1_HEX, '7C3AED'],
        'Figure 2: The five steps from uploading a document to it being ready for AI questions.')
    add_spacer(doc, 6)

    add_h3(doc, 'Steps in plain English')
    steps = [
        ('Step 1 — Upload',         'You select a file and upload it. DocuVault saves it securely.'),
        ('Step 2 — Validate',       'The system checks the file size (must be under 100 MB) and file type.'),
        ('Step 3 — AI Reads It',    'The platform reads the document — extracting all text, tables, and if needed, running text recognition on scanned pages.'),
        ('Step 4 — Save to Index',  'The content is broken into sections and saved in the AI search index. This lets the AI find relevant parts instantly.'),
        ('Step 5 — Ready',          'The document is now available to the AI assistant for any permitted user.'),
    ]
    for title, desc in steps:
        add_body(doc, desc, title + ':')
        add_spacer(doc, 2)

    add_spacer(doc, 6)
    add_info_box(doc,
        'How to index a document',
        'Open any document, then click Index for AI. The status will update to "Indexed" once complete. '
        'Admins can also index all documents at once using the Bulk Index option.')
    add_spacer(doc, 8)

    # 4.3
    add_h2(doc, '4.3  How the AI Answers Questions')
    add_body(doc,
        'When you type a question in the AI assistant, DocuVault goes through these steps '
        'to find and write the answer:')
    add_spacer(doc, 6)

    add_flow_diagram(doc,
        [('User Types\nQuestion', ''),
         ('Platform\nSearches Docs', ''),
         ('AI Reads\nRelevant Parts', ''),
         ('AI Writes\nthe Answer', ''),
         ('User Receives\nAnswer', '')],
        [H2_HEX, ACC_HEX, '059669', H1_HEX, '7C3AED'],
        'Figure 3: How a question becomes an answer.')
    add_spacer(doc, 6)

    add_h3(doc, 'Steps in plain English')
    qsteps = [
        ('Step 1 — You type a question',
         'Ask anything in plain English, e.g. "What are the leave policy rules?"'),
        ('Step 2 — Platform searches your documents',
         'DocuVault searches all indexed documents to find the sections most relevant to your question.'),
        ('Step 3 — AI reads the relevant parts',
         'The AI reads the top matching sections. It only uses documents you have permission to access.'),
        ('Step 4 — AI writes the answer',
         'Using the content it found, the AI writes a clear, direct answer.'),
        ('Step 5 — You receive the answer',
         'The answer appears in the chat with the source documents and page numbers used. You can ask follow-up questions.'),
    ]
    for title, desc in qsteps:
        add_body(doc, desc, title + ':')
        add_spacer(doc, 2)

    add_spacer(doc, 8)
    add_h3(doc, 'The three answer modes')
    add_table(doc, [
        ['Mode',      'How it behaves',                                                         'Best for'],
        ['Hybrid',    'Uses your documents first. Fills gaps with general knowledge if needed.', 'General use — recommended for most teams.'],
        ['Strict',    'Only answers from your documents. Says so clearly if information is not found.', 'Compliance, legal, or regulated content.'],
        ['Indicated', 'Uses both sources but clearly labels which part comes from where.',        'Research, auditing, or when traceability matters.'],
    ], [2.5, 8.0, 6.1])
    doc.add_page_break()


# ── 5. Technical Reference ────────────────────────────────────────────────────

def build_technical(doc):
    add_h1(doc, '5.  Technical Reference')
    add_spacer(doc, 6)

    add_h2(doc, 'Technology Stack')
    add_table(doc, [
        ['Category',          'Technology'],
        ['Web Framework',     'Django 5.x  (Python 3.10+)'],
        ['Database',          'SQLite (development) / PostgreSQL (production)'],
        ['AI Language Model', 'Ollama (local inference) — llama3.1:8b or compatible model\n'
                              'Downloaded once, runs fully offline. No API key required.'],
        ['AI Search Engine',  'ChromaDB vector database — stored locally, no cloud dependency'],
        ['Embedding Model',   'all-MiniLM-L6-v2 — runs on-device via sentence-transformers'],
        ['Document Reading',  'pdfplumber, PyMuPDF, Camelot (tables), Tesseract (OCR for scanned pages)'],
        ['AI Framework',      'LangChain + Ollama backend, HuggingFace (local), PyTorch'],
        ['Max File Upload',   '100 MB per file'],
        ['Search Method',     'Hybrid — 70% semantic similarity + 30% keyword matching'],
        ['Internet Required', 'No — entire stack runs offline on local hardware after setup'],
    ], [5.0, 11.6])
    add_spacer(doc, 10)

    add_h2(doc, 'AI Configuration')
    add_table(doc, [
        ['Setting',                    'Value'],
        ['LLM runtime',                'Ollama — local model server, starts automatically with DocuVault'],
        ['Default model',              'llama3.1:8b  (downloaded once, ~4.7 GB, then fully offline)'],
        ['API key required',           'None — no Groq, OpenAI, or any cloud API key needed'],
        ['Document chunk size',        '512 characters (256 in lightweight mode)'],
        ['Overlap between chunks',     '100 characters'],
        ['Results per query',          '8 chunks (6 in lightweight mode)'],
        ['Conversation memory',        'Up to 8 turns per session'],
        ['Max response length',        '512 tokens (~380 words)'],
        ['Response consistency',       'Low variability — temperature 0.2 for reliable, repeatable answers'],
        ['Internet during inference',  'Not required — model runs entirely on local CPU/GPU'],
    ], [6.0, 10.6])
    add_spacer(doc, 10)

    add_h2(doc, 'User & Permission Model')
    add_table(doc, [
        ['Access Level', 'Who can see the document'],
        ['Public',       'Everyone, including guests who are not logged in.'],
        ['Private',      'The document owner and admins only.'],
        ['Role-Based',   'Users whose role level is equal to or higher than the required level set by the owner.'],
        ['Custom',       'Only specific users selected by the document owner, plus admins.'],
    ], [3.5, 13.1])
    add_spacer(doc, 10)

    add_h2(doc, 'Key Database Records')
    add_table(doc, [
        ['Record Type',        'What it stores'],
        ['Document',           'The file, title, access level, category, tags, and owner.'],
        ['Document Version',   'A snapshot of the document every time it is edited.'],
        ['Chat Session',       'A conversation thread between a user and the AI assistant.'],
        ['Chat Message',       'Each individual question and AI answer, with source references.'],
        ['Activity Log',       'An immutable record of every action in the system.'],
        ['Notification',       'Alerts sent to users for shares, comments, and updates.'],
        ['Shared Link',        'A temporary link with optional password, expiry, and access count.'],
    ], [4.5, 12.1])
    doc.add_page_break()


# ── 6. Add-On Features & Integrations ────────────────────────────────────────

IMG1 = r'D:\AI_Model_Renata\Document-management\Group\V2\DocuVault\clarif_image1.png'
IMG2 = r'D:\AI_Model_Renata\Document-management\Group\V2\DocuVault\clarif_image2.png'


def build_addons(doc):
    add_h1(doc, '6.  Add-On Features & Integrations')
    add_spacer(doc, 6)
    add_body(doc,
        'DocuVault is built to grow with your organisation. The features described in this section '
        'are available as add-ons or integrations that can be configured and enabled on top of the '
        'standard platform — either during initial deployment or at any later stage.')
    add_spacer(doc, 8)

    # 6.1 AMC
    add_h2(doc, '6.1  Annual Maintenance Contract (AMC)')
    add_body(doc,
        'An AMC can be offered to customers to maintain and upgrade the AI components of '
        'DocuVault over time. This ensures the system stays current with the latest AI '
        'models and security standards without any disruption to your data.')
    add_spacer(doc, 4)

    add_h3(doc, 'What the AMC covers')
    amc_bullets = [
        'LLM model upgrades (e.g. upgrading from one AI model to a newer, more capable version)',
        'Performance optimisations and retrieval improvements',
        'Security patches and infrastructure monitoring',
        'UI improvements and new feature rollouts',
    ]
    for b in amc_bullets:
        add_bullet(doc, b)
    add_spacer(doc, 6)

    add_h3(doc, 'Important: No retraining needed for LLM upgrades')
    add_body(doc,
        'DocuVault does NOT train the AI on your documents. It uses a method called '
        'Retrieval-Augmented Generation (RAG). This means:')
    add_bullet(doc, 'LLM upgrades do NOT require your documents to be reprocessed.')
    add_bullet(doc, 'Your existing indexed documents continue working normally after an upgrade.')
    add_bullet(doc,
        'Re-indexing is only needed if the embedding model itself changes '
        '(e.g. switching to a different embedding provider). In that case, '
        'embeddings are recomputed automatically and can be handled under the AMC.')
    add_spacer(doc, 10)

    # 6.2 Email Integration
    add_h2(doc, '6.2  Email System Integration')
    add_body(doc,
        'DocuVault can integrate with your organisation\'s email system. '
        'Emails and their attachments are ingested and indexed so they become '
        'searchable through the AI assistant — just like any other document.')
    add_spacer(doc, 4)

    add_h3(doc, 'Supported email systems')
    email_systems = [
        'Microsoft Outlook / Exchange',
        'Microsoft 365',
        'Gmail',
        'IMAP servers',
        'On-premise enterprise mail servers',
    ]
    for e in email_systems:
        add_bullet(doc, e)
    add_spacer(doc, 6)

    add_h3(doc, 'What gets indexed')
    add_table(doc, [
        ['Email Content',   'What is indexed'],
        ['Email body',      'The full text of the email message.'],
        ['Attachments',     'PDF, Word, Excel, and PowerPoint files attached to emails.'],
        ['Metadata',        'Sender, receiver, subject line, and timestamp.'],
    ], [4.0, 12.6])
    add_spacer(doc, 6)

    add_h3(doc, 'Security & access control')
    add_body(doc,
        'Role-based access ensures that emails are only retrievable by users '
        'who are authorised to see them.')
    add_bullet(doc, 'HR emails → accessible to HR users only')
    add_bullet(doc, 'Finance emails → accessible to Finance users only')
    add_spacer(doc, 4)
    add_body(doc,
        'Once integrated, users can ask the AI assistant questions like: '
        '"Show discussions about vendor contracts with ABC company." '
        'The system retrieves matching content from documents, attachments, '
        'and email conversations together.')
    add_spacer(doc, 8)

    add_h3(doc, 'Email integration flow')
    add_image_centered(doc, IMG2, 9,
        'Figure 4: How emails flow from your mail server into DocuVault and become searchable via the AI assistant.')
    add_spacer(doc, 10)

    # 6.3 LAN Drive Sync
    add_h2(doc, '6.3  LAN Drive Sync (Google Drive-Style)')
    add_body(doc,
        'Instead of manually uploading files through the web interface, users can save '
        'documents into a synchronised folder on their computer. A background sync agent '
        'watches the folder and automatically sends new or updated files to DocuVault.')
    add_spacer(doc, 4)

    add_h3(doc, 'How it works')
    add_flow_diagram(doc,
        [('User PC\n(Sync Folder)', ''),
         ('Sync\nAgent', 'Background'),
         ('DocuVault\nServer', ''),
         ('Version\nControl', ''),
         ('AI\nIndexing', '')],
        [H2_HEX, ACC_HEX, '1D4ED8', '059669', H1_HEX],
        'Figure 5: LAN sync pipeline — from a user\'s folder to DocuVault version control and AI indexing.')
    add_spacer(doc, 6)

    add_h3(doc, 'Features')
    lan_features = [
        'Automatic document synchronisation — no manual uploads required.',
        'Version tracking — every save creates a new version with timestamp and user.',
        'Rollback — restore any previous version from DocuVault.',
        'Instant AI indexing — new documents are indexed automatically as they sync.',
        'Enterprise-ready — supports multiple PCs on the same LAN.',
    ]
    for f in lan_features:
        add_bullet(doc, f)
    add_spacer(doc, 6)

    add_h3(doc, 'Implementation options')
    add_table(doc, [
        ['Option',                          'Description'],
        ['Desktop sync agent (recommended)','A lightweight background app installed on each user\'s PC.'],
        ['Network shared drive monitoring', 'Monitors a shared network folder without any client install.'],
        ['Drive-style desktop client',      'A full Google Drive-like interface with sync status indicators.'],
    ], [5.5, 11.1])
    add_spacer(doc, 10)

    # 6.4 Auto-Indexing
    add_h2(doc, '6.4  Auto-Indexing on Upload')
    add_body(doc,
        'By default, documents must be manually triggered for AI indexing after upload. '
        'The system can be configured to index documents automatically the moment they are '
        'uploaded, so they are immediately available to the AI assistant without any extra step.')
    add_spacer(doc, 4)

    add_h3(doc, 'How indexing works (not AI training)')
    add_body(doc,
        'A common question from clients is whether uploading documents "trains" the AI. '
        'The answer is no. DocuVault uses indexing, not training. The distinction is important:')
    add_table(doc, [
        ['',          'AI Training',                                      'DocuVault Indexing'],
        ['What it is','Modifying the AI model\'s internal weights.',      'Reading and storing document content in a searchable format.'],
        ['Time',      'Hours to days.',                                   'Seconds to minutes per document.'],
        ['Effect',    'Permanently changes the AI model.',                'Adds document to the search index only.'],
        ['Required?', 'Never required in DocuVault.',                     'Required once per document (or on re-upload).'],
    ], [3.5, 7.0, 6.1])
    add_spacer(doc, 6)

    add_h3(doc, 'Indexing pipeline diagram')
    add_image_centered(doc, IMG1, 14,
        'Figure 6: Left — Document processing pipeline (upload to vector database). '
        'Right — Query handling (user question to AI answer).')
    add_spacer(doc, 6)

    add_info_box(doc,
        'How to enable auto-indexing',
        'Auto-indexing can be enabled by your system administrator in the platform configuration. '
        'Once enabled, every document uploaded will be automatically indexed in the background '
        'without any action required from the user.',
        'EFF6FF')
    add_spacer(doc, 14)

    # ── 6.5 Desktop Agent ─────────────────────────────────────────────────────
    add_h2(doc, '6.5  Desktop Agent — Web to Desktop')
    add_body(doc,
        'By default, DocuVault runs entirely in a web browser — no installation needed. '
        'However, for teams who work heavily with local files, a lightweight '
        'Desktop Agent can be installed alongside the web platform. '
        'The agent runs quietly in the background, bridges your computer\'s file system '
        'directly to DocuVault, and makes accessing the platform feel like a native desktop app.')
    add_spacer(doc, 6)

    add_h3(doc, 'What the Desktop Agent does')
    add_table(doc, [
        ['Capability',                    'Description'],
        ['System tray icon',              'The agent sits in your Windows/macOS system tray. One click opens DocuVault instantly — no browser URL to remember.'],
        ['Automatic file watching',       'Select one or more folders on your computer. Any file saved there is automatically detected and synced to DocuVault.'],
        ['Background sync',               'Files are uploaded and indexed in the background without interrupting your work.'],
        ['Auto-connect on startup',       'The agent starts with your computer and reconnects to the DocuVault server automatically — no manual login each day.'],
        ['Local file shortcuts',          'Right-click any file in Windows Explorer / macOS Finder and choose "Send to DocuVault" directly.'],
        ['Offline queue',                 'If the server is temporarily unreachable, the agent queues pending files and syncs them automatically once the connection is restored.'],
        ['Notification alerts',           'Desktop pop-up notifications for completed uploads, indexing status, shared documents, and AI query replies.'],
        ['Session persistence',           'Stay logged in across sessions — the agent maintains your authentication token securely.'],
    ], [5.0, 11.6])
    add_spacer(doc, 8)

    add_h3(doc, 'Web vs Desktop — side by side')
    add_table(doc, [
        ['',                        'Web Browser (Standard)',                    'Desktop Agent (Add-On)'],
        ['Installation',            'None — open any browser',                  'One-time lightweight install (~20 MB)'],
        ['File access',             'Manual upload via browser',                 'Automatic — watches selected folders'],
        ['Startup',                 'Open browser and navigate to URL',          'Starts with computer, always ready'],
        ['Notifications',           'In-browser only',                           'Native desktop pop-ups'],
        ['Right-click upload',      'Not available',                             'Right-click any file to send to DocuVault'],
        ['Offline handling',        'Not available',                             'Queues files and syncs when back online'],
        ['Best for',                'Occasional users, mobile, tablet access',   'Power users who work with many local files daily'],
    ], [3.8, 6.5, 6.3])
    add_spacer(doc, 8)

    add_h3(doc, 'How the Desktop Agent connects')
    add_flow_diagram(doc,
        [('File saved\non your PC', ''),
         ('Agent detects\nthe change', 'Background'),
         ('Uploads to\nDocuVault', 'Auto'),
         ('AI indexes\nthe file', 'Instant'),
         ('Available\neverywhere', 'Web + Desktop')],
        [H2_HEX, ACC_HEX, '1D4ED8', '059669', H1_HEX],
        'Figure 7: From saving a file on your computer to it being searchable in DocuVault — fully automatic.')
    add_spacer(doc, 8)

    add_h3(doc, 'Deployment')
    add_table(doc, [
        ['Supported OS',      'Windows 10/11,  macOS 12+,  Ubuntu 20.04+'],
        ['Server connection', 'Connects to DocuVault over LAN or internet (HTTPS)'],
        ['Authentication',    'Uses the same username / password as the web platform'],
        ['Configuration',     'Admin sets allowed folders, sync rules, and notification preferences via a simple settings panel'],
        ['Distribution',      'Installer provided by Renata AI — deployable via Group Policy (Windows) or MDM (macOS)'],
    ], [4.5, 12.1])
    add_spacer(doc, 8)

    add_info_box(doc,
        'Who should use the Desktop Agent?',
        'The Desktop Agent is ideal for users who regularly save documents locally (reports, drawings, '
        'scanned files, spreadsheets) and want them available in DocuVault without any manual effort. '
        'The web platform remains fully functional alongside it — both modes work together.',
        'EFF6FF')
    add_spacer(doc, 14)

    # ── 6.6 Offline Operation ─────────────────────────────────────────────────
    add_h2(doc, '6.6  Offline Operation — No Internet Required')
    add_body(doc,
        'DocuVault is designed to run entirely on your local network or on a single machine '
        'with no internet connection required at any point during normal operation. '
        'The AI model, the search engine, the database, and the file storage '
        'all run on your own hardware — your data never leaves your premises.')
    add_spacer(doc, 6)

    add_h3(doc, 'What runs locally')
    add_table(doc, [
        ['Component',                  'How it runs offline'],
        ['AI Language Model (LLM)',    'Ollama runs llama3.1:8b (or any compatible model) directly on your server CPU or GPU. '
                                       'Model is downloaded once during setup — no API calls made during use.'],
        ['Embedding Model',            'all-MiniLM-L6-v2 runs on-device via sentence-transformers and PyTorch. '
                                       'No external API needed.'],
        ['Vector Database',            'ChromaDB stores all document embeddings as local files on the server disk.'],
        ['Document Storage',           'All uploaded files are stored in the local file system under /media/.'],
        ['Web Application',            'Django serves the web interface on your local network. '
                                       'Users access it via LAN IP or a local domain — no internet required.'],
        ['Database',                   'SQLite or PostgreSQL runs locally. All user data, roles, and activity logs stay on-site.'],
        ['Desktop Agent',              'Connects to the DocuVault server over LAN (HTTPS). Never calls any external service.'],
    ], [4.5, 12.1])
    add_spacer(doc, 8)

    add_h3(doc, 'Setup — one-time steps (internet needed once only)')
    setup_steps = [
        ('Step 1 — Install DocuVault',
         'Install the platform on your server. This requires internet to download the software package once.'),
        ('Step 2 — Download the AI model',
         'Run: ollama pull llama3.1:8b  (or the model specified by Renata AI). '
         'This downloads the model file (~4.7 GB) to the server. Done once — never again.'),
        ('Step 3 — Download embedding model',
         'The all-MiniLM-L6-v2 embedding model (~90 MB) is downloaded automatically on first run and cached locally.'),
        ('Step 4 — Go fully offline',
         'Once both models are cached, disconnect from the internet. '
         'DocuVault continues to operate with full AI functionality indefinitely.'),
    ]
    for title, desc in setup_steps:
        add_body(doc, desc, title + ':')
        add_spacer(doc, 3)
    add_spacer(doc, 8)

    add_h3(doc, 'Offline capability summary')
    add_table(doc, [
        ['Feature',                        'Works Offline?',  'Notes'],
        ['Document upload & management',   'Yes',             'Files saved locally on server.'],
        ['AI assistant (Q&A)',             'Yes',             'Ollama runs model locally — no API call.'],
        ['Document indexing',              'Yes',             'Embeddings generated on-device.'],
        ['User login & access control',    'Yes',             'Authentication handled by local Django server.'],
        ['Search & filters',               'Yes',             'ChromaDB vector search runs locally.'],
        ['Email integration',              'Yes',             'Connects to on-premise mail server on LAN.'],
        ['Desktop Agent sync',             'Yes',             'Communicates with DocuVault server over LAN.'],
        ['LLM model updates',              'Requires internet once',  'New model pulled via ollama pull, then offline again.'],
        ['Software updates',               'Requires internet once',  'Update package downloaded, then offline again.'],
    ], [5.5, 3.0, 8.1])
    add_spacer(doc, 8)

    add_h3(doc, 'Hardware recommendations for offline deployment')
    add_table(doc, [
        ['Component',   'Minimum',                        'Recommended'],
        ['CPU',         '4 cores  (model runs on CPU)',   '8+ cores for faster AI responses'],
        ['RAM',         '16 GB',                          '32 GB for smooth multi-user operation'],
        ['Storage',     '50 GB free',                     '200 GB+ for large document libraries'],
        ['GPU',         'Not required',                   'NVIDIA GPU (4 GB+ VRAM) for significantly faster AI responses'],
        ['OS',          'Ubuntu 20.04+ / Windows Server', 'Ubuntu 22.04 LTS (recommended)'],
        ['Network',     'LAN only (no internet needed)',  '1 Gbps LAN for fast file sync'],
    ], [3.5, 5.5, 7.6])
    add_spacer(doc, 8)

    add_info_box(doc,
        'Your data stays on your premises — always',
        'Because DocuVault runs entirely on your own hardware with no cloud dependency, '
        'your documents, queries, and AI responses are never transmitted to any external server. '
        'This makes DocuVault suitable for organisations with strict data-sovereignty, '
        'confidentiality, or air-gapped network requirements.',
        OR_HEX)
    doc.add_page_break()


# ── 7. Client Q&A Summary ─────────────────────────────────────────────────────

def build_client_qa(doc):
    add_h1(doc, '7.  Client Q&A Summary')
    add_spacer(doc, 6)
    add_body(doc,
        'The following table summarises the most common questions raised by clients '
        'during technical discussions about DocuVault.')
    add_spacer(doc, 8)

    qa_rows = [
        ['Client Question',                                     'Answer'],
        ['Can you offer an AMC for LLM upgrades?',
         'Yes. An AMC covers LLM upgrades, performance improvements, security patches, and UI enhancements.'],
        ['Do you need to retrain the AI when upgrading the LLM?',
         'No. DocuVault uses RAG — the LLM is not trained on documents. Upgrades do not affect indexed data.'],
        ['When is re-indexing required?',
         'Only if the embedding model itself is changed (e.g. switching providers). This is handled automatically under AMC.'],
        ['Can DocuVault integrate with our email system (Outlook, Exchange)?',
         'Yes. Emails and attachments are ingested and indexed. Users can query email content through the AI assistant.'],
        ['Can it work like Google Drive with automatic sync?',
         'Yes. A LAN Sync Connector can be implemented so documents sync automatically from user folders to DocuVault.'],
        ['Does the AI train automatically when I upload a document?',
         'No. The AI does not train on documents. Uploading triggers indexing, which is fast and does not change the AI model.'],
        ['How is DocuVault different from BigQuery?',
         'BigQuery is designed for structured data analytics using SQL. DocuVault is designed for semantic search and '
         'AI question answering over unstructured documents such as PDFs, emails, and reports.'],
        ['Do users have to open a browser every time?',
         'No. With the optional Desktop Agent, DocuVault runs in the system tray. '
         'Files sync automatically from watched folders, and the platform is accessible with a single click.'],
        ['Can the Desktop Agent work if the server is offline?',
         'Yes. The agent queues pending files locally and syncs them automatically once the server is reachable again.'],
        ['Does the AI require an internet connection or an API key?',
         'No. The AI runs fully offline using Ollama on your local server. '
         'No Groq, OpenAI, or any cloud API key is needed. '
         'The model is downloaded once during setup and then runs without internet.'],
        ['Does data ever leave our network?',
         'Never. All documents, queries, AI responses, and user data stay entirely on your own hardware. '
         'DocuVault has no cloud dependency during operation.'],
        ['What happens if our internet goes down?',
         'Nothing — DocuVault continues to work exactly as normal. '
         'All components (web app, AI model, database, file storage) run on your local server.'],
        ['How much disk space does the AI model need?',
         'The default model (llama3.1:8b) requires approximately 4.7 GB. '
         'This is downloaded once. Additional models can be added without affecting existing operation.'],
    ]
    add_table(doc, qa_rows, [6.0, 10.6])
    add_spacer(doc, 10)

    add_info_box(doc,
        'Have more questions?',
        'If you have additional technical questions not covered in this document, '
        'please contact Renata AI through the channels listed in the Support section.',
        OR_HEX)
    doc.add_page_break()


# ── 8. Support ────────────────────────────────────────────────────────────────

def build_support(doc):
    add_h1(doc, '8.  Support & Contact')
    add_spacer(doc, 8)

    add_body(doc,
        'For assistance, feature requests, or any questions about DocuVault, '
        'please reach out through the following channels:')
    add_spacer(doc, 8)

    add_table(doc, [
        ['Channel',        'Details'],
        ['Website',        'https://renataiot.com'],
        ['Company',        'Renata Envirocom Pvt. Ltd.'],
        ['Product',        'DocuVault v2.0 — Knowledge Management AI Platform'],
        ['Document date',  datetime.date.today().strftime('%B %Y')],
    ], [4.0, 12.6])
    add_spacer(doc, 16)

    add_info_box(doc,
        'Confidentiality Notice',
        'This document is prepared exclusively for authorised DocuVault clients. '
        'The system details and architecture described here are proprietary to Renata Envirocom Pvt. Ltd. '
        'Please do not share or distribute without written consent.',
        OR_HEX)


# ── Main ──────────────────────────────────────────────────────────────────────

def build_docx(output_path: str):
    doc = setup_document()
    add_header_footer(doc)

    build_cover(doc)
    build_toc(doc)
    build_overview(doc)
    build_features(doc)
    build_access(doc)
    build_architecture(doc)
    build_technical(doc)
    build_addons(doc)
    build_client_qa(doc)
    build_support(doc)

    doc.save(output_path)
    print('[OK] DOCX generated: ' + output_path)


if __name__ == '__main__':
    out = r'D:\AI_Model_Renata\Document-management\Group\V2\DocuVault\DocuVault_Client_Reference_Guide.docx'
    build_docx(out)
